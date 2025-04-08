"""
Web Font Analyzer: Scrapes webpages, checks font consistency, and generates AI reports
"""
from http.client import responses

from numpy.distutils.misc_util import clean_up_temporary_directory
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from PIL import Image
import base64
from io import BytesIO
import requests
from docx import Document
import pytesseract
from PIL import Image
import openai

import scrape
import traceback

import tensorflow as tf
from tensorflow.keras.applications.mobilenet import preprocess_input, MobileNet
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Input
import numpy as np
from PIL import Image
import os


from playwright.sync_api import sync_playwright
from transformers import pipeline
from collections import defaultdict
import json
#import re
#import openai
from groq import Groq
import os
import time
from docx import Document
from datetime import datetime
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError

# OPENAI_API_
doc_path = "canvas_analysis.docx"


class FontAnalyzer:
    def __init__(self):
        # Initialize browser and AI components
        self.playwright = sync_playwright().start()
        self.browser = self.playwright.chromium.launch(headless=True)
        self.nlp = pipeline('text2text-generation', model='google/flan-t5-base')

        # Configuration
        self.target_elements = {
            'headings': ['h1', 'h2', 'h3', 'h4', 'h5', 'h6'],
            'paragraphs': ['p'],
            'buttons': ['button', '[type="button"]', '[type="submit"]'],
            'links': ['a']
        }

        # Create directory for canvas screenshots
        self.canvas_output_dir = "canvas_screenshots"
        os.makedirs(self.canvas_output_dir, exist_ok=True)

    def analyze_page(self, url: str) -> dict:
        """Main analysis workflow"""
        page = self.browser.new_page()
        try:
            page.goto(url, wait_until='networkidle')

            # Collect font data
            font_data = self._collect_font_data(page)

            # Analyze consistency
            consistency_report = self._check_font_consistency(font_data)

            # Generate AI report
            ai_analysis = self._generate_ai_report(consistency_report)

            self.check_clickable_elements(url, page)


            self.capture_canvas_screenshots(url)
            self.analyze_screenshots_with_ai()
            scrape.start_scrape(url)

            return {
                'url': url,
                'font_data': font_data,
                'consistency_report': consistency_report,
                'ai_analysis': ai_analysis
            }

        finally:
            page.close()

    def _collect_font_data(self, page) -> dict:
        """Collect font information from page elements"""
        font_data = defaultdict(list)

        # Check all specified elements
        for category, selectors in self.target_elements.items():
            for selector in selectors:
                elements = page.query_selector_all(selector)
                for element in elements:
                    try:
                        styles = element.evaluate('''element => {
                            const style = window.getComputedStyle(element);
                            return {
                                'tag': element.tagName.toLowerCase(),
                                'font_family': style.fontFamily.split(',')[0].replace(/["']/g, '').trim(),
                                'font_size': style.fontSize,
                                'text': element.innerText.slice(0, 50)
                            };
                        }''')
                        font_data[category].append(styles)
                    except:
                        continue

        return dict(font_data)

    def _check_font_consistency(self, font_data: dict) -> dict:
        """Identify font inconsistencies across element categories"""
        report = {}

        for category, elements in font_data.items():
            if not elements:
                continue

            # Get reference values from first element
            ref_font = elements[0]['font_family']
            ref_size = elements[0]['font_size']

            inconsistencies = []
            for idx, element in enumerate(elements[1:]):
                if (element['font_family'] != ref_font or
                        element['font_size'] != ref_size):
                    inconsistencies.append({
                        'element_number': idx + 2,
                        'expected_font': ref_font,
                        'actual_font': element['font_family'],
                        'expected_size': ref_size,
                        'actual_size': element['font_size'],
                        'sample_text': element['text']
                    })

            if inconsistencies:
                report[category] = {
                    'total_elements': len(elements),
                    'inconsistent_count': len(inconsistencies),
                    'reference_font': ref_font,
                    'reference_size': ref_size,
                    'examples': inconsistencies[:3]  # Show first 3 examples
                }

        return report





    from docx import Document
    from playwright.sync_api import TimeoutError as PlaywrightTimeoutError

    def check_clickable_elements(self, url, page):
        """Checks all clickable elements (buttons, links) and logs results to a Word document."""

        output_path = "clickable_element_report.docx"

        # Remove existing report file
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
                print(f"Existing file '{output_path}' deleted.")
            except Exception as e:
                print(f"Error deleting existing file: {e}")
                return

        doc = Document()
        doc.add_heading(f"Clickable Element Report for {url}", level=1)

        # Get clickable elements (safe handles for initial capture only)
        clickable_elements = page.query_selector_all(
            'a, button, [role="button"], [onclick], [type="button"], [type="submit"]')

        print(f"Found {len(clickable_elements)} clickable elements.")

        for i, element in enumerate(clickable_elements):
            desc = "Unnamed element"
            try:
                # Get some description
                try:
                    text = element.inner_text()
                    if text and text.strip():
                        desc = text.strip()
                    else:
                        desc = element.get_attribute('outerHTML')[:100]
                except:
                    desc = element.get_attribute('outerHTML')[:100]

                # Scroll into view
                try:
                    element.scroll_into_view_if_needed(timeout=2000)
                except:
                    pass  # continue even if scrolling fails

                # Try clicking without waiting for navigation
                element.click(timeout=3000, no_wait_after=True, force=True)

                result = f"[{i + 1}] ✅ Click success: {desc}"

            except PlaywrightTimeoutError as te:
                result = f"[{i + 1}] ❌ Click failed: {desc} | Timeout Error: {str(te).splitlines()[0]}"
            except Exception as e:
                result = f"[{i + 1}] ❌ Click failed: {desc} | Reason: {str(e).splitlines()[0]}"

            doc.add_paragraph(result)

        # Save the Word report
        doc.save(output_path)
        print(f"Clickable elements report saved to '{output_path}'")

    @staticmethod
    def get_best_selector(html_snippet):
        """Safely generate a fallback selector from an HTML snippet."""
        if not html_snippet:
            return 'a, button, [role="button"], [onclick], [type="button"], [type="submit"]'

        html_snippet = html_snippet.lower()  # for case-insensitive matching
        if '<a' in html_snippet:
            return 'a'
        if '<button' in html_snippet:
            return 'button'
        if 'type="submit"' in html_snippet:
            return '[type="submit"]'
        if 'type="button"' in html_snippet:
            return '[type="button"]'
        return 'a, button, [role="button"], [onclick], [type="button"], [type="submit"]'

    def _generate_ai_report(self, report: dict) -> str:
        """Generate natural language recommendations using OpenAI's GPT API"""
        if not report:
            return "All fonts are consistent across analyzed elements."

        # Convert dictionary to a string for input
        report_text = json.dumps(report, indent=2)
        if len(report_text) > 1000:  # Truncate if necessary
            report_text = report_text[:1000] + "..."

        print(report_text, '\n')
        prompt = f"""Analyze these web font inconsistencies and provide recommendations:
               {report_text}

               Consider:
               - Design system best practices
               - CSS maintenance strategies
               - User experience impact
               - Technical debt prevention
               Ensure the response is structured in bullet points and actionable suggestions."""

        try:
            client = openai.OpenAI(api_key=OPENAI_API_KEY)

            response = client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": "You are a UI/UX and web design expert. "
                                                  "Study the JSON and give exact and apt remarks about what is wrong and what needs to be done to correct it. "
                                                  "Give output in points and do not give a general roadmap. Be crisp and to the point."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1024
            )

            ai_analysis = response.choices[0].message.content
            return ai_analysis.strip()

        except Exception as e:
            return f"Error generating AI report: {str(e)}"

    """def capture_canvas_screenshots(self, url):
        #Capture screenshots of all canvas elements from a webpage.

        options = Options()
        options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920x1080")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)

        try:
            driver.get(url)
            WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.TAG_NAME, "canvas")))
            time.sleep(2)

            # Execute JavaScript to get Base64 image data from all canvases
            canvas_data = driver.execute_script("""
                     #let canvases = document.querySelectorAll('canvas');
                     #let dataUrls = [];
                     #canvases.forEach((canvas, index) => {
                         #let dataUrl = canvas.toDataURL("image/png");
                         #dataUrls.push({id: canvas.id || `canvas_${index+1}`, dataUrl: dataUrl});
                     #});
                     #return dataUrls;
                 #""")

            #canvas_elements = driver.find_elements(By.TAG_NAME, "canvas")

            #if not canvas_data:
                #print("No canvas elements found on the page.")
                #return

            #print(f"Found {len(canvas_data)} canvas elements. Extracting images...")

            #for canvas in canvas_data:
                #canvas_id = canvas["id"]
                #data_url = canvas["dataUrl"]

                # Remove the header "data:image/png;base64,"
                #image_data = base64.b64decode(data_url.split(",")[1])

                # Save the image
                #img = Image.open(BytesIO(image_data))
                #img.save(os.path.join(self.canvas_output_dir, f"{canvas_id}.png"))

                #print(f"Saved: {canvas_id}.png")

        #except Exception as e:
            #print(f"Error capturing canvas: {e}")

        #finally:
            #driver.quit()"""

    def clean_directory(self):
        """Deletes all files in the specified directory."""
        if os.path.exists(self.canvas_output_dir):
            for file in os.listdir(self.canvas_output_dir):
                file_path = os.path.join(self.canvas_output_dir, file)
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Error deleting {file_path}: {e}")
        else:
            os.makedirs(self.canvas_output_dir)  # Create directory if it doesn't exist

    def capture_canvas_screenshots(self, url):
        """Main method: Sets up the driver and captures canvas + image elements from a webpage."""
        self.clean_directory()

        options = Options()
        options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920x1080")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)

        try:
            driver.get(url)

            # Optional: Scroll down to trigger lazy load
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)

            self.capture_canvas_elements(driver)
            self.capture_image_elements(driver)

        except Exception as e:
            print(f"Error during capture: {e}")
            traceback.print_exc()

        finally:
            driver.quit()

    def capture_canvas_elements(self, driver):
        """Captures all canvas elements and saves them as JPEGs."""
        print("Searching for canvas elements...")

        for i in range(10):
            canvas_elements = driver.find_elements(By.TAG_NAME, "canvas")
            if canvas_elements:
                break
            print(f"Waiting for canvas... ({i + 1}s)")
            time.sleep(1)
        else:
            print("No canvas elements found after waiting.")
            return

        canvas_data = driver.execute_script("""
            let canvases = document.querySelectorAll('canvas');
            let dataUrls = [];
            canvases.forEach((canvas, index) => {
                let dataUrl = canvas.toDataURL("image/jpeg");
                dataUrls.push({id: canvas.id || `canvas_${index+1}`, dataUrl: dataUrl});
            });
            return dataUrls;
        """)

        print(f"Found {len(canvas_data)} canvas elements. Extracting images...")
        for canvas in canvas_data:
            canvas_id = canvas["id"]
            data_url = canvas["dataUrl"]
            image_data = base64.b64decode(data_url.split(",")[1])
            img = Image.open(BytesIO(image_data)).convert("RGB")
            img.save(os.path.join(self.canvas_output_dir, f"{canvas_id}.jpg"), "JPEG")
            print(f"Saved canvas: {canvas_id}.jpg")

    def capture_image_elements(self, driver):
        """Captures all <img> elements and saves them as JPEGs."""
        print("Searching for image elements...")
        img_elements = driver.find_elements(By.TAG_NAME, "img")

        if not img_elements:
            print("No image elements found on the page.")
            return

        print(f"Found {len(img_elements)} image elements. Extracting images...")

        for index, img_element in enumerate(img_elements, start=1):
            try:
                # Try src, then data-src (lazy loading), then srcset
                img_src = (
                        img_element.get_attribute("src") or
                        img_element.get_attribute("data-src") or
                        img_element.get_attribute("srcset")
                )

                if not img_src:
                    print(f"Image #{index} has no valid source.")
                    continue

                # Handle base64-encoded image
                if img_src.startswith("data:image"):
                    header, encoded = img_src.split(",", 1)
                    img_data = base64.b64decode(encoded)
                else:
                    headers = {
                        "User-Agent": "Mozilla/5.0"
                    }
                    response = requests.get(img_src, headers=headers, timeout=10)
                    content_type = response.headers.get("Content-Type", "")
                    if "image" not in content_type:
                        print(f"Image #{index} is not an image: Content-Type={content_type}")
                        continue
                    img_data = response.content

                img = Image.open(BytesIO(img_data)).convert("RGB")
                save_path = os.path.join(self.canvas_output_dir, f"image_{index}.jpg")
                img.save(save_path, "JPEG")
                print(f"Saved image: image_{index}.jpg")

            except Exception as e:
                print(f"Failed to save image #{index}: {e}")



    def analyze_screenshots_with_ai(self):
        """Analyze each canvas screenshot using Groq AI API and save results in a Word document."""

        # Ensure the directory exists
        if not os.path.exists(self.canvas_output_dir):
            print("Canvas screenshots directory does not exist.")
            return

        doc_path = "canvas_analysis.docx"
        if os.path.isfile(doc_path):
            os.remove(doc_path)

        doc = Document()

        # Add table header if new document
        if not os.path.exists(doc_path):
            doc.add_heading("Canvas Image AI Analysis", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Image Name"
            hdr_cells[1].text = "AI Analysis"
        else:
            doc = Document(doc_path)
            table = doc.tables[0]  # Append to existing table

        #groq_api_url = "https://api.groq.com/openai/v1/chat/completions"
       # headers = {"Authorization": f"Bearer {self.groq_key}"}

        #pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

        for filename in sorted(os.listdir(self.canvas_output_dir)):
            if filename.endswith(".jpg"):
                image_path = os.path.join(self.canvas_output_dir, filename)
                image = Image.open(image_path)
                #extracted_text = pytesseract.image_to_string(image)
                #print(extracted_text,"\n")
                # Convert image to base64 for API request
                with open(image_path, "rb") as image_file:
                    base64_image = base64.b64encode(image_file.read()).decode("utf-8")

                # API payload
                """payload = {
                    "model": "llama-3.3-70b-vision",
                    "messages": [
                        {"role": "system",
                         "content": "Read the given image and List the information in the images."},
                        {"role": "user", "content": f"Analyze this UI canvas screenshot:\n\n{base64_image}"}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 1024
                }"""

                try:
                    client = openai.OpenAI(api_key=OPENAI_API_KEY)

                    response = client.chat.completions.create(
                        model="gpt-4-turbo",

                        messages=[
                            {"role": "system", "content": "You are an AI assistant that analyzes images."},
                            {"role": "user", "content": [
                                {"type": "text", "text": "Describe this image."},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                            ]}
                        ],
                        max_tokens=200
                    )

                    # Append AI response to table
                    row_cells = table.add_row().cells
                    row_cells[0].text = filename
                    row_cells[1].text = response.choices[0].message.content
                    #print("\n",response.choices[0].message.content)

                    print(f"Analyzed: {filename}")

                except Exception as e:
                    print(f"Error analyzing {filename}: {e}")
                    continue

        # Save the updated document
        doc.save(doc_path)
        print(f"AI Analysis completed. Results saved in {doc_path}")



def extract_text_from_word(file_path):
        """Extracts text from a Word document."""
        if not os.path.exists(file_path):
            print("File not found!")
            return None

        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return full_text


def chunk_text(text, max_tokens=4000):
    """Splits text into chunks of approximately max_tokens tokens."""
    words = text.split()  # Simple split by words (not perfect tokenization)
    chunks = []
    chunk = []
    token_count = 0

    for word in words:
        token_count += len(word) // 4  # Approximate token estimation (1 token ≈ 4 chars)
        chunk.append(word)

        if token_count >= max_tokens:
            chunks.append(" ".join(chunk))
            chunk = []
            token_count = 0

    if chunk:  # Add any remaining words
        chunks.append(" ".join(chunk))

    return chunks


def intract_openai(text, question):
    """Sends the extracted text in chunks to OpenAI and returns the combined response."""

    chunks = chunk_text(text, max_tokens=4000)  # Adjust chunk size if needed
    responses = []
    client = openai.OpenAI(api_key=OPENAI_API_KEY)

    for chunk in chunks:
        prompt = f"Document Content:\n{chunk}\n\nUser Question: {question}\n\nAnswer:"

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are a helpful assistant that answers questions based on the content of given document. Answer only about the contents."},
                {"role": "user", "content": prompt}
            ]
        )
        responses.append(response.choices[0].message.content)

    # Combine responses into a final answer
    final_response = "\n".join(responses)
    return final_response


def ask_ai(doc):
        #file_path = ""  # Change to your actual file path

        # Extract text from the Word file
        document_text = extract_text_from_word(doc)

        if not document_text:
            print("No text extracted from the file.")
            return

        print("Document loaded successfully! You can now ask questions based on its content.")

        question = input("\nAsk a question (or type 'exit' to quit): ")
        if question.lower() == "exit":
            return 0

        answer = intract_openai(document_text, question)
        print("\nAnswer:", answer)



if __name__ == "__main__":
    analyzer = FontAnalyzer()

    # Example analysis - replace with your target URL
    results = analyzer.analyze_page('https://schoolofscholars.edu.in/blog/stem-education-in-cbse-schools/')

    print("\n=== Font Analysis Report ===")
    print(f"Analyzed URL: {results['url']}")

    # Print consistency findings
    for category, data in results['consistency_report'].items():
        print(f"\n{category.upper()} INCONSISTENCIES:")
        print(f"Elements checked: {data['total_elements']}")
        print(f"Reference font: {data['reference_font']} {data['reference_size']}")
        print(f"Inconsistent elements found: {data['inconsistent_count']}")

    # Print AI analysis
    print("\nAI RECOMMENDATIONS:")
    print(results['ai_analysis'])

    while True:
        print("\nAsk questions")
        print("\nPress 1 for images or 2 for Content of the page")
        resp = input("Enter value\n")
        if resp == "1":
            val = ask_ai(doc_path)
            if val ==0:
                break
        elif resp == "2":
            val = ask_ai(scrape.output_file)
            if val == 0:
                break;
        else:
            print("\nWrong input entred")
            break

