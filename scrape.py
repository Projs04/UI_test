import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os
import pytesseract
import cv2
import numpy as np
from PIL import Image
from io import BytesIO

# Set Tesseract path (Windows users: Update path if needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
output_file = "scraped_data.docx"


def preprocess_image(image):
    """Applies OpenCV preprocessing to improve OCR accuracy."""
    image = np.array(image)  # Convert PIL image to NumPy array

    # Convert grayscale images properly
    if len(image.shape) == 2:  # Already grayscale
        gray = image
    else:
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Convert to grayscale

    blurred = cv2.GaussianBlur(gray, (5, 5), 0)  # Apply Gaussian blur
    _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)  # Binarization
    return Image.fromarray(thresh)  # Convert back to PIL image


def extract_text_from_image(img_url):
    """Download, preprocess, and extract text from an image using OCR."""
    try:
        response = requests.get(img_url)
        image = Image.open(BytesIO(response.content))
        processed_image = preprocess_image(image)  # Apply preprocessing
        text = pytesseract.image_to_string(processed_image)
        return text.strip()
    except Exception as e:
        return f"Error extracting text from image: {e}"

def scrape_website(url):
    """Scrapes all data from a URL and saves it in a structured Word file."""
    response = requests.get(url)
    if response.status_code != 200:
        print("Failed to retrieve the page.")
        return

    soup = BeautifulSoup(response.text, 'html.parser')

    # Create Word document
    doc = Document()
    doc.add_heading(f"Scraped Data from {url}", level=1)

    # Extract headings and paragraphs
    doc.add_heading("Text Content", level=2)
    for tag in ["h1", "h2", "h3", "h4", "h5", "h6", "p"]:
        for element in soup.find_all(tag):
            doc.add_heading(element.get_text(strip=True), level=int(tag[1])) if tag.startswith('h') else doc.add_paragraph(element.get_text(strip=True))

    # Extract links
    doc.add_heading("Extracted Links", level=2)
    for link in soup.find_all('a', href=True):
        doc.add_paragraph(f"{link.get_text(strip=True)}: {link['href']}")

    # Extract tables
    tables = soup.find_all("table")
    if tables:
        doc.add_heading("Extracted Tables", level=2)
        for table in tables:
            table_data = []
            rows = table.find_all("tr")
            for row in rows:
                cols = row.find_all(["td", "th"])
                table_data.append([col.get_text(strip=True) for col in cols])

            # Add table to Word document
            word_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    word_table.cell(i, j).text = cell

    # Extract and process images
    doc.add_heading("Extracted Images and OCR Text", level=2)
    images = soup.find_all("img")
    for img in images:
        img_url = img.get("src")
        if img_url:
            if not img_url.startswith("http"):
                img_url = requests.compat.urljoin(url, img_url)  # Handle relative URLs

            doc.add_paragraph(f"Image URL: {img_url}")
            extracted_text = extract_text_from_image(img_url)
            if extracted_text:
                doc.add_paragraph("Extracted Text:")
                doc.add_paragraph(extracted_text)

    # Save document
    #output_file = "scraped_data.docx"
    doc.save(output_file)
    print(f"Data successfully saved to {output_file}")
    #return output_file

# URL to scrape
#url = "https://www.bits-pilani.ac.in/"  # Replace with the target URL
def start_scrape(url):
    return scrape_website(url)
